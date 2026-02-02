
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Configuration
MARKDOWN_FILE = r'd:\COSMOS\docs\USER_WORKFLOW_GUIDE.md'
OUTPUT_FILE = r'd:\COSMOS\COSMOS_User_Manual.docx'
IMAGE_BASE_DIR = r'd:\COSMOS\docs'  # Images are relative to the markdown file

def create_manual():
    print(f"Reading markdown from: {MARKDOWN_FILE}")
    
    if not os.path.exists(MARKDOWN_FILE):
        print(f"Error: File not found: {MARKDOWN_FILE}")
        return

    with open(MARKDOWN_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title Page
    doc.add_heading('COSMOS PM Admin Panel', 0)
    subtitle = doc.add_paragraph('User Workflow Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # State variables for parsing
    in_table = False
    table_data = []
    in_code_block = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip YAML frontmatter if present (simple check)
        if i == 0 and line == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        # detection of table end
        if in_table:
            if not line.strip().startswith('|'):
                # Process collected table
                process_table(doc, table_data)
                table_data = []
                in_table = False
            else:
                table_data.append(line)
                i += 1
                continue

        # Headers
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            # Clean up links in headers if any [Link](#anchor)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Images: ![Alt](path)
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            # Resolve path
            full_img_path = os.path.join(IMAGE_BASE_DIR, img_path.lstrip('./').replace('/', os.sep))
            
            print(f"Found image: {full_img_path}")
            if os.path.exists(full_img_path):
                try:
                    doc.add_picture(full_img_path, width=Inches(6))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = 'Caption'
                except Exception as e:
                    doc.add_paragraph(f"[Image: {alt_text} - Error inserting image]")
                    print(f"Error inserting image: {e}")
            else:
                doc.add_paragraph(f"[Image: {alt_text} - File not found]")
            
            i += 1
            continue

        # Table start detection
        if line.strip().startswith('|'):
            in_table = True
            table_data.append(line)
            i += 1
            continue

        # Blockquotes / Alerts
        if line.strip().startswith('>'):
            alert_type = "NOTE"
            content = line.lstrip('>').strip()
            
            # Check for GitHub alert syntax > [!NOTE]
            if content.startswith('[!') and ']' in content:
                alert_type_match = re.match(r'\[!(.*?)\]', content)
                if alert_type_match:
                    alert_type = alert_type_match.group(1)
                    # Next lines usually contain the content
                    i += 1
                    content = ""
                    while i < len(lines) and lines[i].strip().startswith('>'):
                        content += lines[i].strip().lstrip('>').strip() + " "
                        i += 1
                    # Add alert box
                    add_alert(doc, alert_type, content)
                    continue
            
            # Standard blockquote
            p = doc.add_paragraph(content)
            p.style = 'Quote'
            i += 1
            continue

        # Lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = parse_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
            
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            text = parse_inline_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Horizontal Rule
        if line.strip() == '---':
            doc.add_paragraph('_' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Standard Paragraph
        if line.strip():
            text = parse_inline_formatting(line)
            doc.add_paragraph(text)
        
        i += 1

    # Final flush if table was last
    if in_table and table_data:
        process_table(doc, table_data)

    doc.save(OUTPUT_FILE)
    print(f"Document saved to {OUTPUT_FILE}")

def parse_inline_formatting(text):
    # Simple bold replacement **text** -> text (would be better to use runs, but for now simple cleanup)
    # Ideally should construct paragraph with runs.
    # For now, let's just strip markdown symbols for cleaner text, 
    # or keep them if we can't easily format inline.
    # A robust solution needs to split text by ** and add runs.
    # For this MVP, we will try to handle bold.
    return text.replace('**', '').replace('__', '')

def process_table(doc, table_lines):
    # Filter out divider lines (---|---|---)
    content_lines = [line for line in table_lines if '---' not in line]
    
    if not content_lines:
        return

    # Determine columns
    header_row = content_lines[0].strip('|').split('|')
    cols = len(header_row)
    rows = len(content_lines)
    
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    for r, line in enumerate(content_lines):
        cells = line.strip().strip('|').split('|')
        row_cells = table.rows[r].cells
        for c, cell_text in enumerate(cells):
            if c < len(row_cells):
                row_cells[c].text = parse_inline_formatting(cell_text.strip())
                
                # Make header bold
                if r == 0:
                    for paragraph in row_cells[c].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    
                    # Add shading to header
                    tcPr = row_cells[c]._tc.get_or_add_tcPr()
                    shd = parse_xml(r'<w:shd {} w:fill="4F81BD"/>'.format(nsdecls('w')))
                    tcPr.append(shd)

def add_alert(doc, alert_type, content):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    
    # Color coding
    color = "E3F2FD" # Blue/Info default
    if "IMPORTANT" in alert_type or "CAUTION" in alert_type:
        color = "FFEBEE" # Red
    elif "WARNING" in alert_type:
        color = "FFF3E0" # Orange
    elif "TIP" in alert_type:
        color = "E8F5E9" # Green
        
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    tcPr.append(shd)
    
    p = cell.paragraphs[0]
    run = p.add_run(f"{alert_type}: ")
    run.bold = True
    p.add_run(content)

if __name__ == '__main__':
    create_manual()

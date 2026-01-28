"""
Script to convert the test case plan markdown to Word document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Create document
doc = Document()

# Set document styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('COSMOS PM Admin Panel - Test Case Plan', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph('This document outlines a comprehensive test case plan for the COSMOS PM Admin Panel, a React-based project management application with Firebase backend. The plan covers all major modules across 5 user roles: SuperAdmin, Admin, Manager, Employee, and Client.')

# Project Overview Table
doc.add_heading('Project Overview', level=1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
data = [
    ('Tech Stack', 'React 19, Vite, TailwindCSS, Firebase (Auth, Firestore)'),
    ('User Roles', 'SuperAdmin, Admin, Manager, Employee (Member), Client'),
    ('Key Modules', 'Dashboard, Resource/Client/Project/Task Management, Lead Management, Calendar, Reports, Documents, Knowledge Base, MOM Generator, Expenses'),
    ('Current Test Status', 'No automated tests exist in the project'),
]
for i, (aspect, details) in enumerate(data):
    row = table.rows[i]
    row.cells[0].text = aspect
    row.cells[1].text = details

doc.add_paragraph()

# Test Categories
def add_test_table(doc, title, tests):
    """Helper to add a test case table"""
    doc.add_heading(title, level=2)
    if not tests:
        return
    
    table = doc.add_table(rows=len(tests) + 1, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    header = table.rows[0]
    headers = ['Test ID', 'Test Case', 'Expected Result', 'Priority']
    for i, text in enumerate(headers):
        cell = header.cells[i]
        cell.text = text
        # Bold header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for i, test in enumerate(tests):
        row = table.rows[i + 1]
        row.cells[0].text = test[0]
        row.cells[1].text = test[1]
        row.cells[2].text = test[2]
        row.cells[3].text = test[3]
    
    doc.add_paragraph()

# Authentication Tests
doc.add_heading('1. Authentication & Authorization Tests', level=1)

login_tests = [
    ('AUTH-001', 'Valid login with correct credentials', 'User redirected to role-specific dashboard', 'High'),
    ('AUTH-002', 'Invalid login with wrong password', 'Error message displayed, user stays on login', 'High'),
    ('AUTH-003', 'Login with non-existent email', 'Appropriate error message shown', 'High'),
    ('AUTH-004', 'Login with empty fields', 'Validation errors shown', 'Medium'),
    ('AUTH-005', 'Session persistence after page refresh', 'User remains logged in', 'High'),
    ('AUTH-006', 'Logout functionality', 'User redirected to login, session cleared', 'High'),
]
add_test_table(doc, '1.1 Login Functionality', login_tests)

password_tests = [
    ('AUTH-007', 'Forgot password with valid email', 'Password reset email sent', 'High'),
    ('AUTH-008', 'Forgot password with invalid email', 'Error message displayed', 'Medium'),
    ('AUTH-009', 'Reset password with valid token', 'Password updated successfully', 'High'),
    ('AUTH-010', 'Force change password on first login', 'User prompted to change password', 'High'),
    ('AUTH-011', 'Change password from profile settings', 'Password updated, user notified', 'Medium'),
]
add_test_table(doc, '1.2 Password Management', password_tests)

rbac_tests = [
    ('AUTH-012', 'SuperAdmin accessing all routes', 'Full access granted', 'High'),
    ('AUTH-013', 'Admin accessing admin-specific routes', 'Access granted to admin routes only', 'High'),
    ('AUTH-014', 'Manager accessing manager routes', 'Access granted to manager routes only', 'High'),
    ('AUTH-015', 'Employee accessing restricted routes', 'Redirected to unauthorized page', 'High'),
    ('AUTH-016', 'Client accessing client-specific routes', 'Access granted to client routes only', 'High'),
    ('AUTH-017', 'Unauthenticated user accessing protected routes', 'Redirected to login', 'High'),
]
add_test_table(doc, '1.3 Role-Based Access Control', rbac_tests)

# Dashboard Tests
doc.add_heading('2. Dashboard Module Tests', level=1)

superadmin_dash = [
    ('DASH-001', 'Dashboard loads with statistics', 'All stat cards display correctly', 'High'),
    ('DASH-002', 'Project statistics accuracy', 'Counts match actual project data', 'High'),
    ('DASH-003', 'Task analytics display', 'Charts render with correct data', 'High'),
    ('DASH-004', 'Resource overview widget', 'Resource list displays correctly', 'Medium'),
    ('DASH-005', 'Recent activity feed', 'Shows recent updates', 'Medium'),
    ('DASH-006', 'Dashboard refresh/reload', 'Data updates in real-time', 'Medium'),
]
add_test_table(doc, '2.1 SuperAdmin Dashboard', superadmin_dash)

role_dash = [
    ('DASH-007', 'Admin dashboard displays admin-specific data', 'Data filtered to admin scope', 'High'),
    ('DASH-008', 'Manager dashboard shows managed resources', 'Team data displayed correctly', 'High'),
    ('DASH-009', 'Employee dashboard shows personal tasks', 'Only assigned tasks visible', 'High'),
    ('DASH-010', 'Client dashboard shows client projects', 'Only client-specific data shown', 'High'),
    ('DASH-011', 'Overdue task count accuracy', 'Correct count of overdue items', 'High'),
]
add_test_table(doc, '2.2 Role-Specific Dashboards', role_dash)

# Resource Management Tests
doc.add_heading('3. Resource Management Tests', level=1)

add_resource = [
    ('RES-001', 'Add new resource with valid data', 'Resource created, appears in list', 'High'),
    ('RES-002', 'Add resource with duplicate email', 'Error message displayed', 'High'),
    ('RES-003', 'Add resource with all optional fields', 'All data saved correctly', 'Medium'),
    ('RES-004', 'Add resource - password toggle functionality', 'Toggle controls password requirements', 'Medium'),
    ('RES-005', 'Form validation for required fields', 'Validation errors shown', 'High'),
]
add_test_table(doc, '3.1 Add Resource', add_resource)

edit_resource = [
    ('RES-006', 'Edit resource basic info', 'Changes saved and reflected', 'High'),
    ('RES-007', 'Edit resource department assignment', 'Department updated correctly', 'Medium'),
    ('RES-008', 'Edit resource skills', 'Skills array updated', 'Medium'),
    ('RES-009', 'Change resource status (active/inactive)', 'Status reflected in UI', 'High'),
]
add_test_table(doc, '3.2 Edit Resource', edit_resource)

view_delete_resource = [
    ('RES-010', 'View resource details modal', 'All information displayed', 'Medium'),
    ('RES-011', 'Delete resource with confirmation', 'Resource removed from system', 'High'),
    ('RES-012', 'Delete resource cancellation', 'Resource remains unchanged', 'Medium'),
    ('RES-013', 'Search resources by name', 'Filtered results displayed', 'Medium'),
    ('RES-014', 'Filter resources by department', 'Correct filtering applied', 'Medium'),
]
add_test_table(doc, '3.3 View & Delete Resource', view_delete_resource)

# Client Management Tests
doc.add_heading('4. Client Management Tests', level=1)

client_tests = [
    ('CLI-001', 'Add new client with valid data', 'Client created successfully', 'High'),
    ('CLI-002', 'Add client with required fields only', 'Client created with defaults', 'Medium'),
    ('CLI-003', 'Edit client company name', 'Name updated in all references', 'High'),
    ('CLI-004', 'Edit client contact details', 'Contact info saved correctly', 'Medium'),
    ('CLI-005', 'Delete client without projects', 'Client removed successfully', 'High'),
    ('CLI-006', 'Delete client with projects', 'Warning/prevention shown', 'High'),
    ('CLI-007', 'View client details', 'All client data displayed', 'Medium'),
    ('CLI-008', 'Search clients', 'Search results accurate', 'Medium'),
]
add_test_table(doc, '4.1 Client CRUD Operations', client_tests)

# Project Management Tests
doc.add_heading('5. Project Management Tests', level=1)

project_crud = [
    ('PROJ-001', 'Create new project with valid data', 'Project created, visible in list', 'High'),
    ('PROJ-002', 'Create project linked to client', 'Client association saved', 'High'),
    ('PROJ-003', 'Edit project name and description', 'Updates saved correctly', 'High'),
    ('PROJ-004', 'Edit project dates', 'Date ranges validated', 'High'),
    ('PROJ-005', 'Delete project', 'Project and relations cleaned up', 'High'),
    ('PROJ-006', 'View project details modal', 'All project info displayed', 'Medium'),
]
add_test_table(doc, '5.1 Project CRUD Operations', project_crud)

project_status = [
    ('PROJ-007', 'Project progress calculation', 'Progress derived from tasks', 'High'),
    ('PROJ-008', 'Seven-stage Kanban view', 'Stages display correctly', 'High'),
    ('PROJ-009', 'Project status change', 'Status updates reflected', 'High'),
    ('PROJ-010', 'Project filtering by status', 'Correct filter results', 'Medium'),
    ('PROJ-011', 'Project search functionality', 'Search works accurately', 'Medium'),
]
add_test_table(doc, '5.2 Project Progress & Status', project_status)

# Task Management Tests
doc.add_heading('6. Task Management Tests', level=1)

task_crud = [
    ('TASK-001', 'Create task with required fields', 'Task created successfully', 'High'),
    ('TASK-002', 'Create task with all fields', 'All data saved correctly', 'Medium'),
    ('TASK-003', 'Create subtask under parent task', 'Subtask linked correctly', 'High'),
    ('TASK-004', 'Edit task details', 'Changes saved and reflected', 'High'),
    ('TASK-005', 'Delete task', 'Task removed, subtasks handled', 'High'),
    ('TASK-006', 'Archive task', 'Task archived, not visible', 'Medium'),
]
add_test_table(doc, '6.1 Task CRUD Operations', task_crud)

kanban_tests = [
    ('TASK-007', 'Kanban board loads with tasks', 'All tasks in correct columns', 'High'),
    ('TASK-008', 'Drag and drop task between columns', 'Task status updated', 'High'),
    ('TASK-009', 'WIP limits enforcement', 'Limits prevent over-assignment', 'Medium'),
    ('TASK-010', 'Kanban filtering by project', 'Correct tasks displayed', 'Medium'),
    ('TASK-011', 'Kanban filtering by assignee', 'Assignee filter works', 'Medium'),
]
add_test_table(doc, '6.2 Kanban Board', kanban_tests)

task_details = [
    ('TASK-012', 'Task modal opens with details', 'All info displayed', 'High'),
    ('TASK-013', 'Task priority change', 'Priority updated', 'Medium'),
    ('TASK-014', 'Task due date change', 'Date validated and saved', 'High'),
    ('TASK-015', 'Task assignee change', 'Assignee updated', 'High'),
    ('TASK-016', 'Task completion with comment', 'Comment modal works', 'Medium'),
    ('TASK-017', 'Time estimate input', 'Estimate saved correctly', 'Low'),
    ('TASK-018', 'Task tags management', 'Tags add/remove works', 'Low'),
]
add_test_table(doc, '6.3 Task Details', task_details)

# Lead Management Tests
doc.add_heading('7. Lead Management Tests', level=1)

lead_crud = [
    ('LEAD-001', 'Add new lead with valid data', 'Lead created successfully', 'High'),
    ('LEAD-002', 'Edit lead details', 'Changes saved correctly', 'High'),
    ('LEAD-003', 'Delete lead', 'Lead removed from system', 'High'),
    ('LEAD-004', 'View lead details modal', 'All lead info displayed', 'Medium'),
    ('LEAD-005', 'Lead grouping/filtering', 'Groups display correctly', 'Medium'),
    ('LEAD-006', 'Lead status change', 'Status updated', 'High'),
]
add_test_table(doc, '7.1 Lead CRUD Operations', lead_crud)

followup_tests = [
    ('LEAD-007', 'Schedule follow-up for lead', 'Follow-up created', 'High'),
    ('LEAD-008', 'Complete follow-up', 'Follow-up marked complete', 'High'),
    ('LEAD-009', 'Reschedule follow-up', 'New date/time saved', 'High'),
    ('LEAD-010', 'Overdue follow-up reminder', 'Reminder popup displays', 'High'),
    ('LEAD-011', 'Follow-up list view', 'List displays correctly', 'Medium'),
    ('LEAD-012', 'Filter follow-ups by status', 'Filter works correctly', 'Medium'),
]
add_test_table(doc, '7.2 Follow-up Management', followup_tests)

lead_settings = [
    ('LEAD-013', 'Add lead source setting', 'Setting saved', 'Medium'),
    ('LEAD-014', 'Add lead status setting', 'Setting saved', 'Medium'),
    ('LEAD-015', 'Edit/Delete settings', 'Settings updated/removed', 'Medium'),
]
add_test_table(doc, '7.3 Lead Settings', lead_settings)

# Calendar Tests
doc.add_heading('8. Calendar Module Tests', level=1)

calendar_display = [
    ('CAL-001', 'Calendar grid loads', 'Current month displayed', 'High'),
    ('CAL-002', 'Navigate to previous/next month', 'Month changes correctly', 'High'),
    ('CAL-003', 'Events display on correct dates', 'Dates aligned properly', 'High'),
    ('CAL-004', 'Meeting requests display', 'Pending requests visible', 'Medium'),
    ('CAL-005', 'Task deadlines on calendar', 'Deadlines marked correctly', 'Medium'),
]
add_test_table(doc, '8.1 Calendar Display', calendar_display)

event_mgmt = [
    ('CAL-006', 'Create new event', 'Event saved and displayed', 'High'),
    ('CAL-007', 'Edit existing event', 'Changes saved', 'High'),
    ('CAL-008', 'Delete event with confirmation', 'Event removed', 'High'),
    ('CAL-009', 'Approve meeting request', 'Request converted to event', 'High'),
    ('CAL-010', 'Cancel/decline meeting', 'Request status updated', 'High'),
]
add_test_table(doc, '8.2 Event Management', event_mgmt)

# Reports Tests
doc.add_heading('9. Reports Module Tests', level=1)

report_tests = [
    ('REP-001', 'Report page loads with default view', 'Initial data displayed', 'High'),
    ('REP-002', 'Project status report generation', 'Accurate data displayed', 'High'),
    ('REP-003', 'Task completion report', 'Data matches actual tasks', 'High'),
    ('REP-004', 'Resource utilization report', 'Utilization calculated correctly', 'Medium'),
    ('REP-005', 'Export report to Excel/CSV', 'File downloads correctly', 'High'),
    ('REP-006', 'Filter reports by date range', 'Date filter works', 'Medium'),
    ('REP-007', 'Filter reports by project/client', 'Filter applied correctly', 'Medium'),
]
add_test_table(doc, '9.1 Reports', report_tests)

# Documents Tests
doc.add_heading('10. Document Management Tests', level=1)

doc_tests = [
    ('DOC-001', 'Upload document', 'File uploaded successfully', 'High'),
    ('DOC-002', 'View document', 'Document viewer opens', 'High'),
    ('DOC-003', 'Download document', 'File downloads correctly', 'High'),
    ('DOC-004', 'Delete document', 'Document removed', 'High'),
    ('DOC-005', 'Document search', 'Search results accurate', 'Medium'),
    ('DOC-006', 'Document access by role', 'Access controlled properly', 'High'),
]
add_test_table(doc, '10.1 Document Operations', doc_tests)

# Knowledge Base Tests
doc.add_heading('11. Knowledge Base Tests', level=1)

kb_tests = [
    ('KB-001', 'Knowledge page loads', 'Content displayed', 'High'),
    ('KB-002', 'View knowledge project detail', 'Details shown correctly', 'High'),
    ('KB-003', 'Add knowledge entry', 'Entry saved', 'Medium'),
    ('KB-004', 'Edit knowledge entry', 'Changes saved', 'Medium'),
    ('KB-005', 'Delete knowledge entry', 'Entry removed', 'Medium'),
    ('KB-006', 'Search knowledge base', 'Search works', 'Medium'),
]
add_test_table(doc, '11.1 Knowledge Base Operations', kb_tests)

# MOM Tests
doc.add_heading('12. MOM (Minutes of Meeting) Generator Tests', level=1)

mom_tests = [
    ('MOM-001', 'Create new MOM', 'MOM saved successfully', 'High'),
    ('MOM-002', 'Add attendees to MOM', 'Attendees saved', 'High'),
    ('MOM-003', 'Add agenda items', 'Agenda items saved', 'High'),
    ('MOM-004', 'Add action items', 'Action items saved', 'High'),
    ('MOM-005', 'Generate PDF export', 'PDF created correctly', 'High'),
    ('MOM-006', 'Edit existing MOM', 'Changes saved', 'Medium'),
    ('MOM-007', 'Delete MOM', 'MOM removed', 'Medium'),
]
add_test_table(doc, '12.1 MOM Operations', mom_tests)

# Expense Tests
doc.add_heading('13. Expense Management Tests', level=1)

expense_tests = [
    ('EXP-001', 'Add new expense', 'Expense created', 'High'),
    ('EXP-002', 'Edit expense details', 'Changes saved', 'High'),
    ('EXP-003', 'Delete expense', 'Expense removed', 'High'),
    ('EXP-004', 'Expense approval workflow', 'Approval status updates', 'High'),
    ('EXP-005', 'Filter expenses by status', 'Filter works', 'Medium'),
    ('EXP-006', 'Expense reports', 'Report generated correctly', 'Medium'),
]
add_test_table(doc, '13.1 Expense Operations', expense_tests)

# Settings Tests
doc.add_heading('14. Settings Module Tests', level=1)

hierarchy_tests = [
    ('SET-001', 'Add department', 'Department created', 'Medium'),
    ('SET-002', 'Edit department', 'Department updated', 'Medium'),
    ('SET-003', 'Delete department', 'Department removed', 'Medium'),
    ('SET-004', 'Add designation', 'Designation created', 'Medium'),
]
add_test_table(doc, '14.1 Hierarchy Settings', hierarchy_tests)

project_settings = [
    ('SET-005', 'Configure project stages', 'Stages saved', 'Medium'),
    ('SET-006', 'Configure task statuses', 'Statuses saved', 'Medium'),
    ('SET-007', 'Status settings visibility', 'Settings apply correctly', 'Medium'),
]
add_test_table(doc, '14.2 Project & Status Settings', project_settings)

theme_tests = [
    ('SET-008', 'Theme toggle (light/dark)', 'Theme changes applied', 'Medium'),
    ('SET-009', 'Profile information update', 'Profile saved', 'Medium'),
    ('SET-010', 'Profile image upload', 'Image saved and displayed', 'Low'),
]
add_test_table(doc, '14.3 Theme & Profile Settings', theme_tests)

# UI/UX Tests
doc.add_heading('15. UI/UX & Cross-Cutting Tests', level=1)

responsive_tests = [
    ('UI-001', 'Desktop viewport (1920x1080)', 'Layout displays correctly', 'High'),
    ('UI-002', 'Tablet viewport (768x1024)', 'Responsive layout', 'Medium'),
    ('UI-003', 'Mobile viewport (375x667)', 'Mobile-friendly layout', 'Medium'),
]
add_test_table(doc, '15.1 Responsive Design', responsive_tests)

modal_tests = [
    ('UI-004', 'Modal open animation', 'Smooth animation', 'Low'),
    ('UI-005', 'Modal close on backdrop click', 'Modal closes', 'Medium'),
    ('UI-006', 'Modal close on Escape key', 'Modal closes', 'Low'),
    ('UI-007', 'Form reset on modal close', 'Form cleared', 'Medium'),
]
add_test_table(doc, '15.2 Modal Interactions', modal_tests)

notification_tests = [
    ('UI-008', 'Success toast notifications', 'Toast appears and auto-dismisses', 'High'),
    ('UI-009', 'Error toast notifications', 'Error displayed clearly', 'High'),
    ('UI-010', 'Loading states (spinners)', 'Spinners show during operations', 'Medium'),
    ('UI-011', 'Skeleton loaders', 'Skeletons during data fetch', 'Medium'),
]
add_test_table(doc, '15.3 Notifications & Feedback', notification_tests)

nav_tests = [
    ('UI-012', 'Sidebar navigation', 'Links work correctly', 'High'),
    ('UI-013', 'Breadcrumb navigation', 'Breadcrumbs accurate', 'Low'),
    ('UI-014', 'Browser back/forward', 'Navigation works', 'Medium'),
]
add_test_table(doc, '15.4 Navigation', nav_tests)

# Real-time Tests
doc.add_heading('16. Real-time Updates Tests', level=1)

realtime_tests = [
    ('RT-001', 'Task update reflects in real-time', 'Other users see update', 'High'),
    ('RT-002', 'New project appears without refresh', 'Real-time subscription works', 'High'),
    ('RT-003', 'Event changes on calendar', 'Calendar updates live', 'Medium'),
    ('RT-004', 'Notification badge updates', 'Badge reflects new items', 'Medium'),
]
add_test_table(doc, '16.1 Real-time Updates', realtime_tests)

# Error Handling Tests
doc.add_heading('17. Error Handling Tests', level=1)

error_tests = [
    ('ERR-001', 'Network error during data fetch', 'User-friendly error shown', 'High'),
    ('ERR-002', 'Firebase offline mode', 'App handles gracefully', 'Medium'),
    ('ERR-003', 'Form submission failure', 'Error message displayed', 'High'),
    ('ERR-004', 'Invalid route access (404)', '404 page or redirect', 'Medium'),
    ('ERR-005', 'Session expiry handling', 'User prompted to re-login', 'High'),
]
add_test_table(doc, '17.1 Error Handling', error_tests)

# Test Count Summary
doc.add_heading('Appendix: Test Count Summary', level=1)

summary_table = doc.add_table(rows=19, cols=5)
summary_table.style = 'Table Grid'

# Header
header = summary_table.rows[0]
for i, text in enumerate(['Module', 'Test Cases', 'High Priority', 'Medium', 'Low']):
    header.cells[i].text = text
    for paragraph in header.cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

summary_data = [
    ('Authentication', '17', '14', '3', '0'),
    ('Dashboard', '11', '8', '3', '0'),
    ('Resource Management', '14', '7', '7', '0'),
    ('Client Management', '8', '4', '4', '0'),
    ('Project Management', '11', '6', '5', '0'),
    ('Task Management', '18', '10', '6', '2'),
    ('Lead Management', '15', '8', '7', '0'),
    ('Calendar', '10', '6', '4', '0'),
    ('Reports', '7', '4', '3', '0'),
    ('Documents', '6', '4', '2', '0'),
    ('Knowledge Base', '6', '2', '4', '0'),
    ('MOM Generator', '7', '4', '3', '0'),
    ('Expenses', '6', '4', '2', '0'),
    ('Settings', '10', '0', '9', '1'),
    ('UI/UX', '14', '4', '7', '3'),
    ('Real-time', '4', '2', '2', '0'),
    ('Error Handling', '5', '3', '2', '0'),
    ('TOTAL', '169', '90', '73', '6'),
]

for i, row_data in enumerate(summary_data):
    row = summary_table.rows[i + 1]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        # Bold the TOTAL row
        if row_data[0] == 'TOTAL':
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

# Footer
doc.add_paragraph()
doc.add_paragraph('Document Version: 1.0')
doc.add_paragraph('Created: January 19, 2026')
doc.add_paragraph('Project: COSMOS PM Admin Panel')

# Save
output_path = r'd:\COSMOS\COSMOS_Test_Case_Plan.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')